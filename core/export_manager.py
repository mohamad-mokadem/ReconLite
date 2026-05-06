import io
from datetime import datetime
from typing import Dict, Any

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ExportManager:
    """Manages export functionality for different report formats"""

    @staticmethod
    def export(scan_data: Dict[str, Any], export_format: str, timestamp: str) -> Dict[str, Any]:
        """Export scan data to the requested format"""
        if export_format == 'pdf':
            buffer = ExportManager.generate_pdf_report(scan_data)
            return {
                'content': buffer.read(),
                'filename': f"recon_report_{scan_data['target']}_{timestamp}.pdf",
                'mimetype': 'application/pdf'
            }
        elif export_format == 'docx':
            buffer = ExportManager.generate_word_report(scan_data)
            return {
                'content': buffer.read(),
                'filename': f"recon_report_{scan_data['target']}_{timestamp}.docx",
                'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
        elif export_format == 'txt':
            text = ExportManager.generate_text_report(scan_data)
            return {
                'content': text.encode(),
                'filename': f"recon_report_{scan_data['target']}_{timestamp}.txt",
                'mimetype': 'text/plain'
            }
        else:
            raise ValueError(f"Unsupported export format: {export_format}")

    @staticmethod
    def generate_pdf_report(scan_data: Dict[str, Any]) -> io.BytesIO:
        """Generate comprehensive PDF report from scan data"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.8 * inch)
        styles = getSampleStyleSheet()
        story = []

        # Title and metadata
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24,
                                     spaceAfter=30, alignment=TA_CENTER, textColor=colors.HexColor('#0066ff'))
        story.append(Paragraph("ReconLite Reconnaissance Report", title_style))
        story.append(Spacer(1, 20))

        # Metadata table
        metadata = [
            ['Target:', scan_data['target']],
            ['Type:', 'Domain' if scan_data['type'] == 'domain' else 'IP Range'],
            ['Scan Time:', scan_data.get('scan_time', 'Unknown')],
            ['Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')]
        ]

        metadata_table = Table(metadata, colWidths=[1.5 * inch, 4 * inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 30))

        # Summary
        section_style = ParagraphStyle('SectionHeader', parent=styles['Heading2'], fontSize=16,
                                       spaceAfter=12, spaceBefore=20, textColor=colors.HexColor('#1f2937'))
        story.append(Paragraph("Executive Summary", section_style))

        summary_data = [['Open Ports:', str(len(scan_data.get('ports_services', [])))]]
        if scan_data['type'] == 'domain':
            dns_count = ExportManager._count_dns_records(scan_data)
            summary_data.extend([
                ['DNS Records:', str(dns_count)],
                ['Subdomains:', str(len(scan_data.get('subdomains', [])))]
            ])

        summary_table = Table(summary_data, colWidths=[2 * inch, 1 * inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f9fafb')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 20))

        # ENHANCED PORTS SECTION with all details
        story.append(Paragraph("Discovered Ports & Services", section_style))
        ports = scan_data.get('ports_services', [])

        if ports:
            # Main ports table with basic info
            ports_data = [['IP Address', 'Port/Protocol', 'Service', 'Version', 'Source']]
            for port in sorted(ports, key=lambda x: x.get('port', 0)):
                ports_data.append([
                    port.get('ip', 'N/A'),
                    f"{port.get('port', 'N/A')}/{port.get('protocol', 'tcp')}",
                    port.get('service', 'unknown'),
                    port.get('version', '')[:25] + ('...' if len(port.get('version', '')) > 25 else ''),
                    port.get('source', 'Unknown')
                ])

            ports_table = Table(ports_data, colWidths=[1.2 * inch, 0.9 * inch, 1.1 * inch, 1.8 * inch, 1 * inch])
            ports_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')])
            ]))
            story.append(ports_table)

            # Additional port details
            story.append(Spacer(1, 15))
            subsection_style = ParagraphStyle('SubsectionHeader', parent=styles['Heading3'], fontSize=14,
                                              spaceAfter=8, spaceBefore=12, textColor=colors.HexColor('#374151'))
            story.append(Paragraph("Port Details", subsection_style))

            for port in sorted(ports, key=lambda x: x.get('port', 0)):
                if port.get('banner') or port.get('http_title') or port.get('asn') or port.get('country'):
                    port_details = []
                    port_details.append([f"Port {port.get('port', 'N/A')} - {port.get('service', 'unknown')}", ""])

                    if port.get('banner'):
                        port_details.append(
                            ['Banner:',
                             port.get('banner', '')[:80] + ('...' if len(port.get('banner', '')) > 80 else '')])
                    if port.get('http_title'):
                        port_details.append(['HTTP Title:', port.get('http_title', '')])
                    if port.get('http_server'):
                        port_details.append(['HTTP Server:', port.get('http_server', '')])
                    if port.get('asn'):
                        port_details.append(['ASN:', port.get('asn', '')])
                    if port.get('country'):
                        port_details.append(['Country:', port.get('country', '')])
                    if port.get('org'):
                        port_details.append(['Organization:', port.get('org', '')])

                    if len(port_details) > 1:  # Only add if there are details beyond the header
                        port_table = Table(port_details, colWidths=[1.5 * inch, 4.5 * inch])
                        port_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
                        ]))
                        story.append(port_table)
                        story.append(Spacer(1, 10))
        else:
            story.append(Paragraph("No open ports discovered", styles['Normal']))

        story.append(Spacer(1, 20))

        # DOMAIN-SPECIFIC SECTIONS
        if scan_data['type'] == 'domain':
            # DNS RECORDS SECTION
            story.append(Paragraph("DNS Intelligence", section_style))

            # SecurityTrails DNS Records
            sec_trails_records = scan_data.get('securitytrails_dns_records', [])
            if sec_trails_records:
                subsection_style = ParagraphStyle('SubsectionHeader', parent=styles['Heading3'], fontSize=14,
                                                  spaceAfter=8, spaceBefore=12, textColor=colors.HexColor('#374151'))
                story.append(Paragraph("SecurityTrails DNS Records", subsection_style))

                # Group by type
                grouped_records = {}
                for record in sec_trails_records:
                    record_type = record.get('type', 'UNKNOWN')
                    if record_type not in grouped_records:
                        grouped_records[record_type] = []
                    grouped_records[record_type].append(record)

                for record_type, records in grouped_records.items():
                    dns_data = [[f'{record_type} Records', 'Value', 'Details']]

                    for record in records:
                        if record_type == 'A':
                            value = record.get('value', '')
                            details = f"Host: {record.get('name', '')}" if record.get('name') else ''
                            if record.get('organization'):
                                details += f"\nOrg: {record.get('organization', '')}"
                        elif record_type == 'MX':
                            value = record.get('name', '')
                            details = f"Priority: {record.get('value', '')}"
                            if record.get('organization'):
                                details += f"\nOrg: {record.get('organization', '')}"
                        else:
                            value = record.get('value', '')
                            details = record.get('name', '') if record.get('name') else ''

                        dns_data.append([record_type, value, details])

                    dns_table = Table(dns_data, colWidths=[1 * inch, 2.5 * inch, 2.5 * inch])
                    dns_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')])
                    ]))
                    story.append(dns_table)
                    story.append(Spacer(1, 10))

            # Regular DNS Records
            regular_dns = scan_data.get('dns_records', {})
            if regular_dns:
                story.append(Paragraph("Standard DNS Records", subsection_style))

                for record_type, values in regular_dns.items():
                    if values:
                        dns_data = [[f'{record_type} Records', 'Value']]
                        for value in values:
                            dns_data.append([record_type, value])

                        dns_table = Table(dns_data, colWidths=[1.5 * inch, 4.5 * inch])
                        dns_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
                        ]))
                        story.append(dns_table)
                        story.append(Spacer(1, 10))

            # WHOIS INFORMATION
            whois_info = scan_data.get('whois_info', {})
            if whois_info and any(whois_info.values()):
                story.append(Spacer(1, 15))
                story.append(Paragraph("Domain Intelligence (Whois)", section_style))

                whois_data = []
                whois_fields = [
                    ('domain_name', 'Domain Name'),
                    ('registrar', 'Registrar'),
                    ('creation_date', 'Creation Date'),
                    ('expiration_date', 'Expiration Date'),
                    ('updated_date', 'Updated Date'),
                    ('domain_age_years', 'Domain Age (Years)'),
                    ('registrant_org', 'Organization'),
                    ('country', 'Country'),
                    ('admin_email', 'Admin Email')
                ]

                for field, label in whois_fields:
                    value = whois_info.get(field, '')
                    if value:
                        whois_data.append([label + ':', str(value)])

                # Name servers
                name_servers = whois_info.get('name_servers', [])
                if name_servers:
                    for i, ns in enumerate(name_servers):
                        label = 'Name Servers:' if i == 0 else ''
                        whois_data.append([label, ns])

                # Status
                status = whois_info.get('status', [])
                if status:
                    for i, s in enumerate(status):
                        label = 'Domain Status:' if i == 0 else ''
                        whois_data.append([label, s])

                if whois_data:
                    whois_table = Table(whois_data, colWidths=[2 * inch, 4 * inch])
                    whois_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
                    ]))
                    story.append(whois_table)

            # SUBDOMAINS
            subdomains = scan_data.get('subdomains', [])
            if subdomains:
                story.append(Spacer(1, 15))
                story.append(Paragraph("Discovered Subdomains", section_style))

                # Create subdomain table with 3 columns
                subdomain_chunks = [subdomains[i:i + 3] for i in range(0, len(subdomains), 3)]
                subdomain_data = []

                for chunk in subdomain_chunks:
                    row = chunk + [''] * (3 - len(chunk))  # Pad with empty strings
                    subdomain_data.append(row)

                subdomain_table = Table(subdomain_data, colWidths=[2 * inch, 2 * inch, 2 * inch])
                subdomain_table.setStyle(TableStyle([
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Courier'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
                    ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f9fafb')])
                ]))
                story.append(subdomain_table)

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_word_report(scan_data: Dict[str, Any]) -> io.BytesIO:
        """Generate comprehensive Word document report from scan data"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")

        doc = Document()

        # Title
        title = doc.add_heading('ReconLite Reconnaissance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_heading('Scan Information', level=1)
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Shading Accent 1'

        metadata_cells = [
            ('Target:', scan_data['target']),
            ('Type:', 'Domain' if scan_data['type'] == 'domain' else 'IP Range'),
            ('Scan Time:', scan_data.get('scan_time', 'Unknown')),
            ('Generated:', datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC'))
        ]

        for i, (label, value) in enumerate(metadata_cells):
            metadata_table.cell(i, 0).text = label
            metadata_table.cell(i, 1).text = value

        # Summary
        doc.add_heading('Executive Summary', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Light Grid Accent 1'

        summary_data = [('Open Ports:', str(len(scan_data.get('ports_services', []))))]
        if scan_data['type'] == 'domain':
            dns_count = ExportManager._count_dns_records(scan_data)
            summary_data.extend([
                ('DNS Records:', str(dns_count)),
                ('Subdomains:', str(len(scan_data.get('subdomains', []))))
            ])

        for label, value in summary_data:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        # ENHANCED PORTS SECTION
        doc.add_heading('Discovered Ports & Services', level=1)
        ports = scan_data.get('ports_services', [])

        if ports:
            # Main ports table
            ports_table = doc.add_table(rows=1, cols=6)
            ports_table.style = 'Light Grid Accent 1'

            hdr_cells = ports_table.rows[0].cells
            hdr_cells[0].text = 'IP Address'
            hdr_cells[1].text = 'Port'
            hdr_cells[2].text = 'Service'
            hdr_cells[3].text = 'Version'
            hdr_cells[4].text = 'Country'
            hdr_cells[5].text = 'Source'

            for port in sorted(ports, key=lambda x: x.get('port', 0)):
                row_cells = ports_table.add_row().cells
                row_cells[0].text = port.get('ip', 'N/A')
                row_cells[1].text = f"{port.get('port', 'N/A')}/{port.get('protocol', 'tcp')}"
                row_cells[2].text = port.get('service', 'unknown')
                row_cells[3].text = port.get('version', '')[:30]
                row_cells[4].text = port.get('country', '')
                row_cells[5].text = port.get('source', 'Unknown')

            # Port details
            doc.add_heading('Port Details', level=2)
            for port in sorted(ports, key=lambda x: x.get('port', 0)):
                if port.get('banner') or port.get('http_title') or port.get('asn'):
                    doc.add_heading(f"Port {port.get('port', 'N/A')} - {port.get('service', 'unknown')}", level=3)

                    details = []
                    if port.get('banner'):
                        details.append(f"Banner: {port.get('banner', '')}")
                    if port.get('http_title'):
                        details.append(f"HTTP Title: {port.get('http_title', '')}")
                    if port.get('http_server'):
                        details.append(f"HTTP Server: {port.get('http_server', '')}")
                    if port.get('asn'):
                        details.append(f"ASN: {port.get('asn', '')}")
                    if port.get('org'):
                        details.append(f"Organization: {port.get('org', '')}")

                    for detail in details:
                        para = doc.add_paragraph(detail)
                        para.style = 'List Bullet'
        else:
            doc.add_paragraph('No open ports discovered')

        # DOMAIN-SPECIFIC SECTIONS
        if scan_data['type'] == 'domain':
            doc.add_page_break()

            # DNS RECORDS
            doc.add_heading('DNS Intelligence', level=1)

            # SecurityTrails DNS Records
            sec_trails_records = scan_data.get('securitytrails_dns_records', [])
            if sec_trails_records:
                doc.add_heading('SecurityTrails DNS Records', level=2)

                # Group by type
                grouped_records = {}
                for record in sec_trails_records:
                    record_type = record.get('type', 'UNKNOWN')
                    if record_type not in grouped_records:
                        grouped_records[record_type] = []
                    grouped_records[record_type].append(record)

                for record_type, records in grouped_records.items():
                    doc.add_heading(f'{record_type} Records', level=3)
                    dns_table = doc.add_table(rows=1, cols=3)
                    dns_table.style = 'Light Grid Accent 2'

                    hdr_cells = dns_table.rows[0].cells
                    hdr_cells[0].text = 'Name/Host'
                    hdr_cells[1].text = 'Value'
                    hdr_cells[2].text = 'Organization'

                    for record in records:
                        row_cells = dns_table.add_row().cells
                        row_cells[0].text = record.get('name', '')
                        row_cells[1].text = record.get('value', '')
                        row_cells[2].text = record.get('organization', '')

            # Regular DNS Records
            regular_dns = scan_data.get('dns_records', {})
            if regular_dns:
                doc.add_heading('Standard DNS Records', level=2)

                for record_type, values in regular_dns.items():
                    if values:
                        doc.add_heading(f'{record_type} Records', level=3)
                        for value in values:
                            doc.add_paragraph(value, style='List Bullet')

            # WHOIS INFORMATION
            whois_info = scan_data.get('whois_info', {})
            if whois_info and any(whois_info.values()):
                doc.add_heading('Domain Intelligence (Whois)', level=1)

                whois_table = doc.add_table(rows=1, cols=2)
                whois_table.style = 'Light Shading Accent 2'

                whois_fields = [
                    ('domain_name', 'Domain Name'),
                    ('registrar', 'Registrar'),
                    ('creation_date', 'Creation Date'),
                    ('expiration_date', 'Expiration Date'),
                    ('updated_date', 'Updated Date'),
                    ('domain_age_years', 'Domain Age (Years)'),
                    ('registrant_org', 'Organization'),
                    ('country', 'Country'),
                    ('admin_email', 'Admin Email')
                ]

                for field, label in whois_fields:
                    value = whois_info.get(field, '')
                    if value:
                        row_cells = whois_table.add_row().cells
                        row_cells[0].text = label + ':'
                        row_cells[1].text = str(value)

                # Name servers
                name_servers = whois_info.get('name_servers', [])
                if name_servers:
                    doc.add_heading('Name Servers', level=2)
                    for ns in name_servers:
                        doc.add_paragraph(ns, style='List Bullet')

                # Status
                status = whois_info.get('status', [])
                if status:
                    doc.add_heading('Domain Status', level=2)
                    for s in status:
                        doc.add_paragraph(s, style='List Bullet')

            # SUBDOMAINS
            subdomains = scan_data.get('subdomains', [])
            if subdomains:
                doc.add_heading('Discovered Subdomains', level=1)

                # Add subdomains in a structured way
                subdomain_para = doc.add_paragraph()
                for i, subdomain in enumerate(sorted(subdomains)):
                    if i > 0:
                        subdomain_para.add_run(' • ')
                    subdomain_para.add_run(subdomain)
                    if (i + 1) % 4 == 0:  # New line every 4 subdomains
                        subdomain_para.add_run('\n')

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_text_report(scan_data: Dict[str, Any]) -> str:
        """Generate comprehensive plain text report from scan data"""
        report = f"""ReconLite Scan Report
========================
Target: {scan_data['target']}
Type: {'Domain' if scan_data['type'] == 'domain' else 'IP Range'}
Scan Time: {scan_data.get('scan_time', 'Unknown')}

Summary
-------
Open Ports: {len(scan_data.get('ports_services', []))}
"""

        if scan_data['type'] == 'domain':
            dns_count = ExportManager._count_dns_records(scan_data)
            report += f"""DNS Records: {dns_count}
Subdomains: {len(scan_data.get('subdomains', []))}

"""

        # Ports section
        report += 'Discovered Ports & Services\n---------------------------\n'
        ports = scan_data.get('ports_services', [])
        if ports:
            for port in sorted(ports, key=lambda x: x.get('port', 0)):
                report += f"{port.get('ip', 'N/A')}\t{port.get('port', 'N/A')}/{port.get('protocol', 'tcp')}\t{port.get('service', 'unknown')}"
                if port.get('version'):
                    report += f" {port.get('version')}"
                report += '\n'

            # Add detailed port information
            report += '\nPort Details\n------------\n'
            for port in sorted(ports, key=lambda x: x.get('port', 0)):
                if port.get('banner') or port.get('http_title') or port.get('asn') or port.get('country'):
                    report += f"\nPort {port.get('port', 'N/A')} - {port.get('service', 'unknown')}:\n"
                    if port.get('banner'):
                        report += f"  Banner: {port.get('banner', '')}\n"
                    if port.get('http_title'):
                        report += f"  HTTP Title: {port.get('http_title', '')}\n"
                    if port.get('http_server'):
                        report += f"  HTTP Server: {port.get('http_server', '')}\n"
                    if port.get('asn'):
                        report += f"  ASN: {port.get('asn', '')}\n"
                    if port.get('country'):
                        report += f"  Country: {port.get('country', '')}\n"
                    if port.get('org'):
                        report += f"  Organization: {port.get('org', '')}\n"
            report += '\n'
        else:
            report += 'No open ports discovered\n\n'

        # Domain-specific sections
        if scan_data['type'] == 'domain':
            # DNS Records
            report += 'DNS Records\n-----------\n'
            dns_records = scan_data.get('dns_records', {})
            securitytrails_records = scan_data.get('securitytrails_dns_records', [])

            if securitytrails_records:
                report += 'SecurityTrails DNS Records:\n'
                grouped_records = {}
                for record in securitytrails_records:
                    record_type = record.get('type', 'UNKNOWN')
                    if record_type not in grouped_records:
                        grouped_records[record_type] = []
                    grouped_records[record_type].append(record)

                for record_type, records in grouped_records.items():
                    report += f'{record_type} Records:\n'
                    for record in records:
                        if record_type == 'A':
                            report += f"  {record.get('value', '')} ({record.get('name', '')})\n"
                        elif record_type == 'MX':
                            report += f"  {record.get('name', '')} (Priority: {record.get('value', '')})\n"
                        else:
                            report += f"  {record.get('value', '')}\n"
                    report += '\n'

            if dns_records:
                report += 'Standard DNS Records:\n'
                for record_type, values in dns_records.items():
                    if values:
                        report += f'{record_type} Records:\n'
                        for value in values:
                            report += f"  {value}\n"
                        report += '\n'

            if not securitytrails_records and not dns_records:
                report += 'No DNS records found\n\n'

            # Whois Information
            whois_info = scan_data.get('whois_info', {})
            if whois_info and any(whois_info.values()):
                report += 'Whois Information\n-----------------\n'
                whois_fields = [
                    ('domain_name', 'Domain Name'),
                    ('registrar', 'Registrar'),
                    ('creation_date', 'Creation Date'),
                    ('expiration_date', 'Expiration Date'),
                    ('updated_date', 'Updated Date'),
                    ('domain_age_years', 'Domain Age (Years)'),
                    ('registrant_org', 'Organization'),
                    ('country', 'Country'),
                    ('admin_email', 'Admin Email')
                ]

                for field, label in whois_fields:
                    value = whois_info.get(field, '')
                    if value:
                        report += f'{label}: {value}\n'

                # Name servers
                name_servers = whois_info.get('name_servers', [])
                if name_servers:
                    report += '\nName Servers:\n'
                    for ns in name_servers:
                        report += f"  {ns}\n"

                # Status
                status = whois_info.get('status', [])
                if status:
                    report += '\nDomain Status:\n'
                    for s in status:
                        report += f"  {s}\n"

                report += '\n'
            else:
                report += 'Whois Information\n-----------------\nNo whois information available\n\n'

            # Subdomains
            subdomains = scan_data.get('subdomains', [])
            report += 'Discovered Subdomains\n--------------------\n'
            if subdomains:
                for subdomain in sorted(subdomains):
                    report += f"{subdomain}\n"
            else:
                report += 'No subdomains discovered\n'

        return report

    @staticmethod
    def _count_dns_records(scan_data: Dict[str, Any]) -> int:
        """Count total DNS records from all sources"""
        count = 0

        # Count regular DNS records
        dns_records = scan_data.get('dns_records', {})
        if dns_records:
            count += len([k for k, v in dns_records.items() if v and len(v) > 0])

        # Count SecurityTrails DNS records
        securitytrails_records = scan_data.get('securitytrails_dns_records', [])
        if securitytrails_records:
            unique_types = set(record.get('type', '') for record in securitytrails_records)
            count += len(unique_types)

        return count